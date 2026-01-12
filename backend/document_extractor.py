"""
Document text extraction service
Extracts text from various document formats (PDF, DOCX, TXT, etc.)
"""
import os
import logging
from typing import Optional

# PDF extraction
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("pypdf not available. PDF extraction will not work.")

# DOCX extraction
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX extraction will not work.")


class DocumentExtractor:
    """Extract text from various document formats"""
    
    @staticmethod
    def extract_text_from_file(file_path: str) -> Optional[str]:
        """
        Extract text from a file based on its extension
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text or None if extraction failed
        """
        if not os.path.exists(file_path):
            logging.error(f"File not found: {file_path}")
            return None
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        try:
            if ext == '.pdf':
                return DocumentExtractor._extract_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return DocumentExtractor._extract_from_docx(file_path)
            elif ext in ['.txt', '.md', '.py', '.java', '.cpp', '.c', '.js', '.html', '.css', '.json', '.xml']:
                return DocumentExtractor._extract_from_text(file_path)
            else:
                logging.warning(f"Unsupported file format: {ext}")
                return None
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> Optional[str]:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            logging.error("pypdf not available. Cannot extract from PDF.")
            return None
        
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            full_text = '\n'.join(text_parts)
            return full_text.strip() if full_text else None
        except Exception as e:
            logging.error(f"Error extracting PDF text: {str(e)}")
            return None
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> Optional[str]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            logging.error("python-docx not available. Cannot extract from DOCX.")
            return None
        
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = '\n'.join(text_parts)
            return full_text.strip() if full_text else None
        except Exception as e:
            logging.error(f"Error extracting DOCX text: {str(e)}")
            return None
    
    @staticmethod
    def _extract_from_text(file_path: str) -> Optional[str]:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                return text.strip() if text else None
            except UnicodeDecodeError:
                # Fallback to latin-1 if UTF-8 fails
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return text.strip() if text else None
        except Exception as e:
            logging.error(f"Error extracting text file: {str(e)}")
            return None
    
    @staticmethod
    def get_text_preview(text: str, max_length: int = 500) -> str:
        """
        Get a preview of the extracted text
        
        Args:
            text: Full text
            max_length: Maximum length of preview
            
        Returns:
            Preview text
        """
        if not text:
            return ""
        
        if len(text) <= max_length:
            return text
        
        return text[:max_length] + "..."

